"""
Report and Resume Generation Service for CandiLift
"""

import os
import tempfile
from datetime import datetime
from typing import Dict, List, Any
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


class ReportGenerator:
    """Generates PDF analysis reports"""
    
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self.setup_custom_styles()
    
    def setup_custom_styles(self):
        """Setup custom paragraph styles"""
        # Title style
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1,  # Center
            textColor=colors.HexColor('#1e40af')
        ))
        
        # Section header style
        self.styles.add(ParagraphStyle(
            name='SectionHeader',
            parent=self.styles['Heading2'],
            fontSize=16,
            spaceAfter=12,
            spaceBefore=20,
            textColor=colors.HexColor('#374151')
        ))
        
        # Score style
        self.styles.add(ParagraphStyle(
            name='ScoreStyle',
            parent=self.styles['Normal'],
            fontSize=14,
            alignment=1,  # Center
            textColor=colors.HexColor('#059669')
        ))
    
    def generate_pdf_report(self, analysis_result: Dict[str, Any], filename: str = None) -> str:
        """Generate PDF analysis report"""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"candilift_report_{timestamp}.pdf"
        
        # Create temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
            temp_path = temp_file.name
        
        # Create PDF document
        doc = SimpleDocTemplate(temp_path, pagesize=A4)
        story = []
        
        # Title
        story.append(Paragraph("CandiLift Resume Analysis Report", self.styles['CustomTitle']))
        story.append(Spacer(1, 20))
        
        # Analysis date
        story.append(Paragraph(f"Analysis Date: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", 
                              self.styles['Normal']))
        story.append(Spacer(1, 30))
        
        # Overall Scores
        story.append(Paragraph("Overall Scores", self.styles['SectionHeader']))
        
        scores_data = [
            ['Score Type', 'Score', 'Status'],
            ['ATS Compatibility', f"{analysis_result.get('ats_score', 0)}%", 
             self._get_score_status(analysis_result.get('ats_score', 0))],
            ['Recruiter Appeal', f"{analysis_result.get('recruiter_score', 0)}%", 
             self._get_score_status(analysis_result.get('recruiter_score', 0))],
            ['Likelihood of Hearing Back', f"{analysis_result.get('likelihood_score', 0)}%", 
             self._get_score_status(analysis_result.get('likelihood_score', 0))]
        ]
        
        scores_table = Table(scores_data, colWidths=[2*inch, 1.5*inch, 1.5*inch])
        scores_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f3f4f6')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#374151')),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(scores_table)
        
        # Add likelihood explanation if available
        if analysis_result.get('likelihood_explanation'):
            story.append(Spacer(1, 15))
            story.append(Paragraph("Likelihood Assessment", self.styles['SectionHeader']))
            story.append(Paragraph(analysis_result['likelihood_explanation'], self.styles['Normal']))
            story.append(Spacer(1, 8))
            story.append(Paragraph("Reality Check: Most job postings receive 100-500+ applications. Only 2-5% typically get any response.", 
                                  self.styles['Normal']))
        
        story.append(Spacer(1, 30))
        
        # Score Drivers
        if 'score_drivers' in analysis_result:
            story.append(Paragraph("Score Breakdown", self.styles['SectionHeader']))
            
            for driver in analysis_result['score_drivers']:
                story.append(Paragraph(f"• {driver['component']}: {driver['score']}%", 
                                      self.styles['Normal']))
                story.append(Paragraph(f"  {driver['explanation']}", 
                                      self.styles['Normal']))
                story.append(Spacer(1, 8))
        
        story.append(Spacer(1, 20))
        
        # Recommendations
        if 'recommendations' in analysis_result:
            story.append(Paragraph("Recommendations", self.styles['SectionHeader']))
            
            for i, rec in enumerate(analysis_result['recommendations'], 1):
                story.append(Paragraph(f"{i}. {rec['category']}", 
                                      self.styles['Normal']))
                story.append(Paragraph(f"   {rec['description']}", 
                                      self.styles['Normal']))
                story.append(Paragraph(f"   Estimated Improvement: +{rec['estimated_lift']}%", 
                                      self.styles['Normal']))
                if 'example' in rec:
                    story.append(Paragraph(f"   Example: {rec['example']}", 
                                          self.styles['Normal']))
                story.append(Spacer(1, 12))
        
        story.append(Spacer(1, 20))
        
        # Gap Analysis
        if 'gap_analysis' in analysis_result:
            story.append(Paragraph("Skills Analysis", self.styles['SectionHeader']))
            
            gap = analysis_result['gap_analysis']
            
            # Present Skills
            if gap.get('present_skills'):
                story.append(Paragraph("✓ Skills Found in Resume:", self.styles['Normal']))
                for skill in gap['present_skills'][:10]:  # Limit to 10
                    story.append(Paragraph(f"  • {skill}", self.styles['Normal']))
                story.append(Spacer(1, 12))
            
            # Missing Skills
            if gap.get('missing_skills'):
                story.append(Paragraph("⚠ Missing Skills:", self.styles['Normal']))
                for skill in gap['missing_skills'][:10]:  # Limit to 10
                    story.append(Paragraph(f"  • {skill}", self.styles['Normal']))
        
        story.append(Spacer(1, 30))
        
        # Footer
        story.append(Paragraph("Generated by CandiLift - AI-Powered Resume Analysis", 
                              self.styles['Normal']))
        
        # Build PDF
        doc.build(story)
        
        return temp_path
    
    def _get_score_status(self, score: int) -> str:
        """Get status text for score"""
        if score >= 80:
            return "Excellent"
        elif score >= 60:
            return "Good"
        elif score >= 40:
            return "Fair"
        else:
            return "Needs Improvement"


class ResumeGenerator:
    """Generates improved resume DOCX files"""
    
    def __init__(self):
        pass
    
    def generate_improved_resume(self, original_resume_text: str, 
                                analysis_result: Dict[str, Any], 
                                filename: str = None) -> str:
        """Generate improved resume based on analysis"""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"improved_resume_{timestamp}.docx"
        
        # Create temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            temp_path = temp_file.name
        
        # Create new document
        doc = Document()
        
        # Set up styles
        self._setup_styles(doc)
        
        # Add header with contact info (if available)
        self._add_header(doc, original_resume_text)
        
        # Add professional summary (if missing)
        if analysis_result.get('recommendations'):
            self._add_professional_summary(doc, analysis_result)
        
        # Add skills section with improvements
        self._add_skills_section(doc, analysis_result)
        
        # Add experience section
        self._add_experience_section(doc, original_resume_text)
        
        # Add education section
        self._add_education_section(doc, original_resume_text)
        
        # Add improvements based on recommendations
        self._add_improvements(doc, analysis_result)
        
        # Save document
        doc.save(temp_path)
        
        return temp_path
    
    def _setup_styles(self, doc: Document):
        """Setup document styles"""
        # Title style
        title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Calibri'
        title_style.font.size = Inches(0.2)
        title_style.font.bold = True
        
        # Heading style
        heading_style = doc.styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
        heading_style.font.name = 'Calibri'
        heading_style.font.size = Inches(0.12)
        heading_style.font.bold = True
        
        # Normal style
        normal_style = doc.styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
        normal_style.font.name = 'Calibri'
        normal_style.font.size = Inches(0.11)
    
    def _add_header(self, doc: Document, resume_text: str):
        """Add header section"""
        # Extract name (first line usually)
        lines = resume_text.split('\n')
        name = lines[0].strip() if lines else "Your Name"
        
        # Add name
        name_para = doc.add_paragraph(name)
        name_para.style = 'CustomTitle'
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add contact info placeholder
        contact_para = doc.add_paragraph("Email: your.email@example.com | Phone: (555) 123-4567 | LinkedIn: linkedin.com/in/yourprofile")
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacing
    
    def _add_professional_summary(self, doc: Document, analysis_result: Dict[str, Any]):
        """Add professional summary section"""
        doc.add_paragraph("PROFESSIONAL SUMMARY", style='CustomHeading')
        
        # Generate summary based on analysis
        ats_score = analysis_result.get('ats_score', 0)
        recruiter_score = analysis_result.get('recruiter_score', 0)
        
        if ats_score >= 70 and recruiter_score >= 70:
            summary_text = "Experienced professional with strong technical skills and proven track record of delivering results. Seeking to leverage expertise in a challenging role."
        else:
            summary_text = "Motivated professional with diverse skills and experience. Committed to continuous learning and contributing to team success."
        
        doc.add_paragraph(summary_text, style='CustomNormal')
        doc.add_paragraph()  # Spacing
    
    def _add_skills_section(self, doc: Document, analysis_result: Dict[str, Any]):
        """Add skills section with improvements"""
        doc.add_paragraph("TECHNICAL SKILLS", style='CustomHeading')
        
        # Get skills from analysis
        gap_analysis = analysis_result.get('gap_analysis', {})
        present_skills = gap_analysis.get('present_skills', [])
        missing_skills = gap_analysis.get('missing_skills', [])
        
        # Add present skills
        if present_skills:
            skills_text = "Programming Languages: " + ", ".join(present_skills[:5])
            doc.add_paragraph(skills_text, style='CustomNormal')
        
        # Add missing skills as suggestions
        if missing_skills:
            suggestions_text = "Consider Adding: " + ", ".join(missing_skills[:3])
            doc.add_paragraph(suggestions_text, style='CustomNormal')
        
        doc.add_paragraph()  # Spacing
    
    def _add_experience_section(self, doc: Document, resume_text: str):
        """Add experience section"""
        doc.add_paragraph("PROFESSIONAL EXPERIENCE", style='CustomHeading')
        
        # Extract experience from original text
        lines = resume_text.split('\n')
        experience_lines = [line for line in lines if any(keyword in line.lower() 
                           for keyword in ['experience', 'employment', 'work', 'position'])]
        
        if experience_lines:
            for line in experience_lines[:3]:  # Limit to 3 entries
                doc.add_paragraph(line.strip(), style='CustomNormal')
        else:
            doc.add_paragraph("• [Add your work experience here]", style='CustomNormal')
            doc.add_paragraph("• Include job titles, companies, and key achievements", style='CustomNormal')
        
        doc.add_paragraph()  # Spacing
    
    def _add_education_section(self, doc: Document, resume_text: str):
        """Add education section"""
        doc.add_paragraph("EDUCATION", style='CustomHeading')
        
        # Extract education from original text
        lines = resume_text.split('\n')
        education_lines = [line for line in lines if any(keyword in line.lower() 
                          for keyword in ['education', 'university', 'college', 'degree', 'bachelor', 'master'])]
        
        if education_lines:
            for line in education_lines[:2]:  # Limit to 2 entries
                doc.add_paragraph(line.strip(), style='CustomNormal')
        else:
            doc.add_paragraph("• [Add your education details here]", style='CustomNormal')
        
        doc.add_paragraph()  # Spacing
    
    def _add_improvements(self, doc: Document, analysis_result: Dict[str, Any]):
        """Add improvement suggestions"""
        recommendations = analysis_result.get('recommendations', [])
        
        if recommendations:
            doc.add_paragraph("IMPROVEMENT SUGGESTIONS", style='CustomHeading')
            
            for i, rec in enumerate(recommendations[:3], 1):  # Limit to 3
                doc.add_paragraph(f"{i}. {rec['description']}", style='CustomNormal')
                if 'example' in rec:
                    doc.add_paragraph(f"   Example: {rec['example']}", style='CustomNormal')
            
            doc.add_paragraph()  # Spacing


class FileService:
    """Handles file operations and cleanup"""
    
    @staticmethod
    def cleanup_temp_file(file_path: str):
        """Clean up temporary file"""
        try:
            if os.path.exists(file_path):
                os.unlink(file_path)
        except Exception as e:
            print(f"Error cleaning up temp file {file_path}: {e}")
    
    @staticmethod
    def get_file_content(file_path: str) -> bytes:
        """Read file content as bytes"""
        try:
            with open(file_path, 'rb') as f:
                return f.read()
        except Exception as e:
            print(f"Error reading file {file_path}: {e}")
            return b""
